import docx
from docx import Document
import sys,os , logging
from colorama import init, Fore, Style
import pandas as pd
from pprint import pprint


# Initialize colorama
init()

# Function to redirect console output to a log file
def redirect_console_output(log_file):
    # Create a logger
    logger = logging.getLogger('console_logger')
    logger.setLevel(logging.DEBUG)

    # Create a file handler and set the log level
    file_handler = logging.FileHandler(log_file)
    file_handler.setLevel(logging.DEBUG)

    # Create a log formatter
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')

    # Add the formatter to the file handler
    file_handler.setFormatter(formatter)

    # Add the file handler to the logger
    logger.addHandler(file_handler)

    # Redirect console output to the logger
    sys.stdout = LoggerWriter(logger, logging.log())


# Custom class to redirect console output to the logger
class LoggerWriter:
    def __init__(self, logger, level):
        self.logger = logger
        self.level = level

    def write(self, message):
        if message != '\n':
            self.logger.log(self.level, message)

            # Add color and style to console output
            if self.level == logging.INFO:
                colored_message = f"{Fore.GREEN}{Style.BRIGHT}{message}{Style.RESET_ALL}"
            elif self.level == logging.ERROR:
                colored_message = f"{Fore.RED}{Style.BRIGHT}{message}{Style.RESET_ALL}"
            else:
                colored_message = message

            sys.__stdout__.write(colored_message)

    def flush(self):
        pass

# Function to export log file to Word document
def export_to_word(log_file, output_file):
    # Open the log file
    with open(log_file, 'r') as file:
        log_data = file.read()

    # Create a new Word document
    doc = docx.Document()

    # Add the log file content to the document
    doc.add_paragraph(log_data)

    # Save the document
    doc.save(output_file)

def read_file(file_path: str) -> pd.DataFrame:
    # Derive the file extension using os module

    SUPPORTED_EXTENSIONS = {
        ".csv": pd.read_csv,
        ".xlsx": pd.read_excel,
        ".xls": pd.read_excel,
        ".json": pd.read_json,
        ".parquet": pd.read_parquet,
        ".pickle": pd.read_pickle,
        ".feather": pd.read_feather,
        ".html": pd.read_html,
        ".sql": pd.read_sql,
        ".txt": pd.read_table
    }
    READ_FUNCTIONS = {
        ".csv": pd.read_csv,
        ".xlsx": pd.read_excel,
        ".xls": pd.read_excel,
        ".json": pd.read_json,
        ".parquet": pd.read_parquet,
        ".pickle": pd.read_pickle,
        ".feather": pd.read_feather,
        ".html": pd.read_html,
        ".sql": pd.read_sql,
        ".txt": pd.read_table
    }
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    # Ensure the file extension is supported
    assert file_extension in SUPPORTED_EXTENSIONS, "Unsupported file format. Only CSV, Excel, JSON, Parquet, Pickle, Feather, HTML, SQL, and Text files are supported."

    # Retrieve the appropriate read function based on the file extension
    read_function = READ_FUNCTIONS[file_extension]

    # Read the file using the read function
    df = read_function(file_path)

    return df

def validate_dataframe(df, required_cols):
    """
    Validates the dataframe for required columns and emptiness.
    """
    assert isinstance(df, pd.DataFrame), "Input must be a pandas DataFrame"

    # Check if required columns are present
    required_cols = [col.strip().lower() for col in required_cols]
    missing_cols = [col for col in required_cols if col not in [col.strip().lower() for col in df.columns]]

    if missing_cols:
        print("Missing columns in dataframe:")
        pprint(missing_cols, indent=4)
        return False

    # Check if the dataframe is empty
    if df.empty:
        print("Dataframe is empty.")
        return False

    return True

def generate_summary_report(total_rows, process_rows, no_process_rows, golden_rows, hr_review_rows):
    # Create a new Word document
    doc = Document()

    # Add title to document
    doc.add_heading('Data Summary Statistics', 0)

    # Add summary statistics to document
    doc.add_paragraph(f"Total rows: {total_rows}")
    doc.add_paragraph(f"Rows with Process Step: {process_rows} ({process_rows / total_rows * 100:.2f}%)")
    doc.add_paragraph(f"Rows without Process Step: {no_process_rows} ({no_process_rows / total_rows * 100:.2f}%)")
    doc.add_paragraph(f"Rows in golden source: {golden_rows} ({golden_rows / total_rows * 100:.2f}%)")
    doc.add_paragraph(f"Rows for HR manual review: {hr_review_rows} ({hr_review_rows / total_rows * 100:.2f}%)")

    # Save the document
    doc.save('data_summary.docx')